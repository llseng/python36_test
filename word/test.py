# -*- coding: utf-8 -*-
# @Descripttion: 
# @version: 
# @Author: llseng
# @Date: 2021-01-22 18:18:49
# @LastEditors: llseng
# @LastEditTime: 2021-01-22 18:18:50
#

from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading( "文档标题", 0 )

p = document.add_paragraph( "这是段落, 这是段落, 这是段落, 这是段落" )
p.add_run( "bold" ).bold = True
p.add_run( " and some " )
p.add_run( "italic" ).italic = True

document.add_heading( "Heading, level 1", 1 )
document.add_paragraph( "Intense quote", style='Intense Quote' )

#添加项目列表（前面一个小圆点）
document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph('second item in unordered list', style='List Bullet')

#添加项目列表（前面数字）
document.add_paragraph('first item in ordered list', style='List Number')
document.add_paragraph('second item in ordered list', style='List Number')

document.add_picture( 'source/bt.jpg', width=Inches( 2 ) )

cash_list = (
    (1,'0.3元',3000,0.3),
    (2,'0.3元',3000,0.3),
    (3,'1元',10000,1),
    (4,'2元',20000,2),
    (5,'3元',30000,3),
    (6,'5元',50000,5),
    (7,'10元',100000,10),
    (8,'50元',500000,50),
    (9,'100元',1000000,100)
)

cash_table = document.add_table( rows=1, cols=4, style='Light Shading Accent 3' )

cash_table.rows[0].cells[0].text = 'id'
cash_table.rows[0].cells[1].text = '简介'
cash_table.rows[0].cells[2].text = '积分'
cash_table.rows[0].cells[3].text = '金额'

for cid, cdesc, cinte, csum in cash_list:
    row_cols = cash_table.add_row().cells
    row_cols[0].text = str( cid )
    row_cols[1].text = str( cdesc )
    row_cols[2].text = str( cinte )
    row_cols[3].text = str( csum )

document.add_page_break()

document.save('logs/demo.docx')